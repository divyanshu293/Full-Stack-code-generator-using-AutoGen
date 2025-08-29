import os
from docx import Document
from PyPDF2 import PdfReader
import markdown
from exceptions.custom_exceptions import ParsingError, UnsupportedFileFormatError


def parse_file(file_path):
    """
    Parse a file and return its raw text and structured sections.
    """
    if not os.path.exists(file_path):
        raise ParsingError("File does not exist")
 
    file_extension = os.path.splitext(file_path)[1].lower()
 
    try:
        if file_extension == '.pdf':
            return parse_pdf(file_path)
        elif file_extension == '.docx':
            return parse_docx(file_path)
        elif file_extension == '.md':
            return parse_markdown(file_path)
        else:
            raise UnsupportedFileFormatError(f"Unsupported file format: {file_extension}")
    except Exception as e:
        raise ParsingError(f"Error parsing file: {e}")
 
def parse_pdf(file_path):
    """
    Parse a PDF file and return its raw text and structured sections.
    """
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text, {"sections": text.split('\n\n')}
    except Exception as e:
        raise ParsingError(f"Error parsing PDF: {e}")
 
def parse_docx(file_path):
    """
    Parse a DOCX file and return its raw text and structured sections.
    """
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text, {"sections": [para.text for para in doc.paragraphs]}
    except Exception as e:
        raise ParsingError(f"Error parsing DOCX: {e}")
 
def parse_markdown(file_path):
    """
    Parse a Markdown file and return its raw text and structured sections.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            html = markdown.markdown(text)
        return text, {"sections": html.split('<h1>')[1:]}  # Split by top-level headers
    except Exception as e:
        raise ParsingError(f"Error parsing Markdown: {e}")
